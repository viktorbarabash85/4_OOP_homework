"""
Расширенное дополнение к домашнему заданию 14_1_homework
c учетом изменений домашнего задания 15_1
по взаимодействию пользователя с данными JSON
"""

# 15_1_homework

import json
from pathlib import Path
from typing import List

from openpyxl.workbook import Workbook

from src.category import Category
from src.product import Product


def load_products_from_json(filepath: str = "data/products.json") -> List[Category]:
    """
    Загружает товары и категории из JSON-файла и создает соответствующие объекты.

    :param filepath: Путь к JSON-файлу. Если путь относительный, считается относительно корня проекта.
    :return: Список объектов Category, созданных из данных JSON.
    :raises FileNotFoundError: Если файл не найден.
    """
    p = Path(filepath)
    if not p.is_absolute():
        # Если путь относительный, считаем его относительно корня проекта (две родительские директории от этого файла)
        p = Path(__file__).resolve().parent.parent / filepath
    if not p.exists():
        raise FileNotFoundError(f"Файл {p} не найден.")

    with p.open(encoding="utf-8") as file:
        data = json.load(file)

    # Если данные представлены в виде словаря с ключом "categories", используем его значение.
    if isinstance(data, dict) and "categories" in data:
        data = data["categories"]

    categories: List[Category] = []
    for cat_data in data:
        products = []
        for prod_data in cat_data.get("products", []):
            product = Product(
                prod_data.get("name"), prod_data.get("description"), prod_data.get("price"), prod_data.get("quantity")
            )
            products.append(product)
        category = Category(cat_data.get("name"), cat_data.get("description"), products)
        categories.append(category)
    return categories


def save_products_to_json(categories: List[Category], filepath: str = "data/products_.json") -> None:
    """
    Сохраняет список категорий и товаров в JSON-файл.

    :param categories: Список объектов Category.
    :param filepath: Путь для сохранения файла.
    """
    p = Path(filepath)
    if not p.is_absolute():
        p = Path(__file__).resolve().parent.parent / filepath

    data = []
    for category in categories:
        data.append(
            {
                "name": category.name,
                "description": category.description,
                # Используем метод get_products_list() для получения списка объектов Product
                "products": [
                    {
                        "name": prod.name,
                        "description": prod.description,
                        "price": prod.price,
                        "quantity": prod.quantity,
                    }
                    for prod in category.get_products_list()
                ],
            }
        )

    with p.open("w", encoding="utf-8") as file:
        json.dump(data, file, indent=4, ensure_ascii=False)


def export_products_to_file(
    categories: List[Category], output_format: str = "txt", base_filename: str = "products"
) -> None:
    """
    Экспортирует список категорий и товаров в указанный формат: txt, xlsx или docx.
    Файл сохраняется в папку "data" в корневой директории проекта.

    :param categories: Список объектов Category.
    :param output_format: Формат экспорта ("txt", "xlsx", "docx").
    :param base_filename: Базовое имя файла (без расширения).
    """
    output_format = output_format.lower()
    # Определяем базовую директорию для сохранения файлов
    base_dir = Path(__file__).resolve().parent.parent / "data"
    base_dir.mkdir(exist_ok=True)  # Создаем директорию, если она не существует

    if output_format == "txt":
        export_filename = base_dir / f"{base_filename}_.txt"
        content = ""
        for cat in categories:
            content += f"Категория: {cat.name}\n"
            content += f"Описание: {cat.description}\n"
            content += "Товары:\n"
            # Используем метод get_products_list() для получения списка объектов Product
            for product in cat.get_products_list():
                content += f"{str(product)}\n"
            content += "\n" + "=" * 50 + "\n"
        with export_filename.open("w", encoding="utf-8") as f:
            f.write(content)

    elif output_format == "xlsx":
        export_filename = base_dir / f"{base_filename}_.xlsx"
        wb = Workbook()  # Создаем рабочую книгу
        ws = wb.active  # Получаем активный рабочий лист
        if ws is None:
            raise RuntimeError("Активный лист не найден!")
        ws.title = "Продукты"
        # Заголовки таблицы
        ws.append(["Категория", "Описание категории", "Название продукта", "Описание продукта", "Цена", "Количество"])
        for cat in categories:
            for prod in cat.get_products_list():
                ws.append([cat.name, cat.description, prod.name, prod.description, prod.price, prod.quantity])
        wb.save(export_filename)

    elif output_format == "docx":
        try:
            from docx import Document
        except ImportError:
            print("Для экспорта в docx установите пакет python-docx.")
            return
        export_filename = base_dir / f"{base_filename}_.docx"
        document = Document()
        for cat in categories:
            document.add_heading(f"Категория: {cat.name}", level=1)
            document.add_paragraph(f"Описание: {cat.description}")
            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Название продукта"
            hdr_cells[1].text = "Описание продукта"
            hdr_cells[2].text = "Цена"
            hdr_cells[3].text = "Количество"
            for prod in cat.get_products_list():
                row_cells = table.add_row().cells
                row_cells[0].text = prod.name
                row_cells[1].text = prod.description
                row_cells[2].text = str(prod.price)
                row_cells[3].text = str(prod.quantity)
            document.add_paragraph("")  # Разделитель между категориями
        document.save(export_filename)
    else:
        print(f"Неподдерживаемый формат: {output_format}")
